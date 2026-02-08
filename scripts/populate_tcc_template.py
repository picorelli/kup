#!/usr/bin/env python3
"""
Populates template-tcc.docx with content from tcc-reescrita.md.

Formatting applied to match ABNT / USP-Esalq TCC norms:
- Body paragraphs: Arial 11, 1.5 line spacing (w:line=360), first-line indent 1.25 cm (w:firstLine=709)
- Sub-headings:    Same spacing + indent, bold
- Captions:        Simple spacing (w:line=240), no indent
- Table body:      Simple spacing, no bold, col 1 left-aligned, others centered
- Section headings: inherited from List Paragraph style (unchanged)

Output: doc/tcc-final.docx
"""

import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Low-level pPr builders ───────────────────────────────────────────────────

def _make_pPr(style="Normal", spacing_line="360", first_line=None, jc=None):
    """
    Build a w:pPr element.
    spacing_line: "360" = 1.5x, "240" = single
    first_line:   twips string, e.g. "709" for 1.25 cm
    jc:           None (default) or "both" (justify), "left", "center"
    """
    pPr = OxmlElement("w:pPr")

    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style)
    pPr.append(pStyle)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), spacing_line)
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    pPr.append(spacing)

    if first_line is not None:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), first_line)
        pPr.append(ind)

    if jc is not None:
        jc_elem = OxmlElement("w:jc")
        jc_elem.set(qn("w:val"), jc)
        pPr.append(jc_elem)

    return pPr


def _make_run(text, bold=False, italic=False):
    """Build a w:r element."""
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    return run


def _para_after(ref_para, pPr, runs):
    """Insert a new paragraph immediately after ref_para. Returns a Paragraph proxy."""
    new_p = OxmlElement("w:p")
    ref_para._element.addnext(new_p)
    new_p.append(pPr)
    for run in runs:
        new_p.append(run)

    from docx.text.paragraph import Paragraph as DocxParagraph
    return DocxParagraph(new_p, ref_para._parent)


# ── Public helpers ────────────────────────────────────────────────────────────

def insert_body_para(ref_para, text, bold=False, italic=False):
    """Body paragraph: 1.5 spacing, 1.25 cm first-line indent."""
    pPr = _make_pPr(spacing_line="360", first_line="709")
    runs = [_make_run(text, bold=bold, italic=italic)] if text else []
    return _para_after(ref_para, pPr, runs)


def insert_body_para_mixed(ref_para, runs_data):
    """Body paragraph with mixed bold/italic runs."""
    pPr = _make_pPr(spacing_line="360", first_line="709")
    runs = [_make_run(t, bold=b, italic=i) for t, b, i in runs_data if t]
    return _para_after(ref_para, pPr, runs)


def insert_caption_para(ref_para, text, bold=False):
    """Caption (Tabela N, Fonte:): simple spacing, no indent."""
    pPr = _make_pPr(spacing_line="240", first_line=None)
    runs = [_make_run(text, bold=bold)] if text else []
    return _para_after(ref_para, pPr, runs)


def insert_placeholder_para(ref_para, text):
    """Placeholder text in italic, body spacing."""
    pPr = _make_pPr(spacing_line="360", first_line="709")
    runs = [_make_run(f"[{text}]", italic=True)]
    return _para_after(ref_para, pPr, runs)


def insert_blank_para(ref_para):
    """Blank paragraph (used as cursor after table insertion)."""
    pPr = _make_pPr(spacing_line="360", first_line=None)
    return _para_after(ref_para, pPr, [])


def remove_paragraph(para):
    """Remove a paragraph from the document."""
    para._element.getparent().remove(para._element)


def get_paragraphs_between(doc, start_para, end_texts):
    """
    Return paragraphs between start_para (exclusive) and
    the first paragraph whose text matches end_texts (exclusive).
    Uses XML element identity to avoid proxy object mismatch.
    """
    found = False
    result = []
    start_elem = start_para._element
    for p in doc.paragraphs:
        if p._element is start_elem:
            found = True
            continue
        if found:
            if p.text.strip() in end_texts:
                break
            result.append(p)
    return result


def find_heading_para(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


# ── Table insertion ───────────────────────────────────────────────────────────

def insert_table_after(ref_para, headers, rows):
    """
    Insert a table after ref_para following ABNT table rules:
    - No bold anywhere (not even header row)
    - Col 1: left-aligned; cols 2+: center-aligned
    - Simple spacing in cells
    Returns the table XML element.
    """
    body = ref_para._element.getparent()
    ref_index = list(body).index(ref_para._element)

    tbl = OxmlElement("w:tbl")

    # Table properties
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tbl.append(tblPr)

    # tblGrid
    all_rows = [headers] + rows
    num_cols = max(len(r) for r in all_rows) if all_rows else 1
    tblGrid = OxmlElement("w:tblGrid")
    for _ in range(num_cols):
        tblGrid.append(OxmlElement("w:gridCol"))
    tbl.append(tblGrid)

    # Rows
    for row_data in all_rows:
        tr = OxmlElement("w:tr")
        for col_idx, cell_text in enumerate(row_data):
            tc = OxmlElement("w:tc")
            # Cell paragraph
            p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            # Alignment: col 0 → left, others → center
            if col_idx > 0:
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "center")
                pPr.append(jc)
            # Simple spacing in table cells
            spacing = OxmlElement("w:spacing")
            spacing.set(qn("w:line"), "240")
            spacing.set(qn("w:lineRule"), "auto")
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            pPr.append(spacing)
            p.append(pPr)
            # Run (no bold per ABNT rule)
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            r.append(rPr)
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = cell_text
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)

    body.insert(ref_index + 1, tbl)
    return tbl


# ── Markdown Parser ───────────────────────────────────────────────────────────

def parse_inline(text):
    """Split on *italic* markers → list of (text, bold, italic) tuples."""
    runs = []
    parts = re.split(r'\*(.*?)\*', text)
    for i, part in enumerate(parts):
        if part:
            runs.append((part, False, i % 2 == 1))
    return runs or [(text, False, False)]


def parse_md_section(lines):
    """
    Parse markdown lines into block dicts:
      {'type': 'para', 'runs': [(text, bold, italic), ...]}
      {'type': 'subheading', 'text': str}
      {'type': 'table', 'headers': [...], 'rows': [[...]]}
      {'type': 'caption', 'text': str}
      {'type': 'placeholder', 'text': str}
    """
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line or line == '---':
            i += 1
            continue

        if line.startswith('### '):
            blocks.append({'type': 'subheading', 'text': line[4:].strip()})
            i += 1
            continue

        if line.startswith('## ') or line.startswith('# '):
            i += 1
            continue

        if line.startswith('|'):
            tlines = []
            while i < len(lines) and lines[i].rstrip().startswith('|'):
                tlines.append(lines[i].rstrip())
                i += 1
            if len(tlines) >= 2:
                headers = [c.strip() for c in tlines[0].split('|') if c.strip()]
                rows = [[c.strip() for c in tl.split('|') if c.strip()]
                        for tl in tlines[2:] if tl.strip()]
                blocks.append({'type': 'table', 'headers': headers, 'rows': rows})
            continue

        if line.startswith('Tabela ') or line.startswith('Fonte:'):
            blocks.append({'type': 'caption', 'text': line})
            i += 1
            continue

        if line.startswith('*[') and line.endswith(']*'):
            blocks.append({'type': 'placeholder', 'text': line[2:-2]})
            i += 1
            continue

        # Regular paragraph — accumulate continuation lines
        para_lines = [line]
        i += 1
        while i < len(lines):
            nl = lines[i].rstrip()
            if not nl or nl.startswith('#') or nl.startswith('|') or nl == '---':
                break
            para_lines.append(nl)
            i += 1

        full_text = ' '.join(para_lines)
        blocks.append({'type': 'para', 'runs': parse_inline(full_text)})

    return blocks


def parse_tcc_reescrita(path):
    """Parse tcc-reescrita.md, merging Resultados/Discussão into one section."""
    with open(path, encoding='utf-8') as f:
        content = f.read()

    parts = re.split(r'^## (.+)$', content, flags=re.MULTILINE)
    sections = {}
    for idx in range(1, len(parts), 2):
        heading = parts[idx].strip()
        body = parts[idx + 1] if idx + 1 < len(parts) else ''
        sections[heading] = parse_md_section(body.split('\n'))

    # Merge result/discussion subsections
    merged = []
    for key in ['Resultados Preliminares', 'Resultados Finais', 'Discussão']:
        if key in sections:
            merged.append({'type': 'subheading', 'text': key})
            merged.extend(sections.pop(key))
    if merged:
        sections['Resultados e Discussão'] = merged

    return sections


# ── Template Population ───────────────────────────────────────────────────────

SECTION_HEADING_MAP = {
    'Introdução':                        'Introdução',
    'Metodologia ou Material e Métodos': 'Metodologia ou Material e Métodos',
    'Resultados e Discussão':            'Resultados e Discussão',
    'Considerações Finais':              'Conclusão(ões) ou Considerações Finais',
    'Referências':                       'Referências',
}

NEXT_HEADINGS = {
    'Introdução':                        ['Metodologia ou Material e Métodos'],
    'Metodologia ou Material e Métodos': ['Resultados e Discussão'],
    'Resultados e Discussão':            ['Conclusão(ões) ou Considerações Finais'],
    'Considerações Finais':              ['Agradecimento (opcional, 1 parágrafo, bem sucinto)'],
    'Referências':                       ['Apêndice ou Anexo (opcional)'],
}


def populate_section(doc, md_heading, blocks):
    template_heading = SECTION_HEADING_MAP[md_heading]
    stop_texts = NEXT_HEADINGS[md_heading]

    heading_para = find_heading_para(doc, template_heading)
    if heading_para is None:
        print(f"  WARNING: heading not found: {template_heading!r}")
        return

    # Delete instruction paragraphs
    to_delete = get_paragraphs_between(doc, heading_para, stop_texts)
    for p in to_delete:
        remove_paragraph(p)

    # Insert content blocks (each appended after cursor, advancing it)
    cursor = heading_para
    for block in blocks:
        t = block['type']
        if t == 'subheading':
            cursor = insert_body_para(cursor, block['text'], bold=True)

        elif t == 'para':
            runs = block['runs']
            has_markup = any(b or it for _, b, it in runs)
            if has_markup:
                cursor = insert_body_para_mixed(cursor, runs)
            else:
                cursor = insert_body_para(cursor, runs[0][0] if runs else '')

        elif t == 'placeholder':
            cursor = insert_placeholder_para(cursor, block['text'])

        elif t == 'caption':
            cursor = insert_caption_para(cursor, block['text'])

        elif t == 'table':
            tbl = insert_table_after(cursor, block['headers'], block['rows'])
            # Advance cursor past the table with a blank para
            cursor = insert_blank_para(cursor)
            # The blank para was inserted BEFORE the table (addnext on cursor).
            # We need to move it AFTER the table. Do this by re-inserting.
            body = cursor._element.getparent()
            blank_elem = cursor._element
            tbl_index = list(body).index(tbl)
            body.remove(blank_elem)
            body.insert(tbl_index + 1, blank_elem)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    md_path = 'doc/tcc-reescrita.md'
    template_path = 'doc/template-tcc.docx'
    output_path = 'doc/tcc-final.docx'

    print("Parsing tcc-reescrita.md...")
    sections = parse_tcc_reescrita(md_path)
    for heading, blocks in sections.items():
        print(f"  '{heading}': {len(blocks)} blocks")

    print("\nLoading template...")
    doc = Document(template_path)

    for md_heading in SECTION_HEADING_MAP:
        if md_heading not in sections:
            print(f"  SKIP (not in MD): {md_heading}")
            continue
        print(f"Populating: {SECTION_HEADING_MAP[md_heading]}")
        populate_section(doc, md_heading, sections[md_heading])

    print(f"\nSaving to {output_path}...")
    doc.save(output_path)
    print("Done.")


if __name__ == '__main__':
    main()
